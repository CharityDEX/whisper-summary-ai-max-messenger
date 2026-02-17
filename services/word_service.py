"""
Word document creation service for transcript formatting.
Based on the PDF service logic but adapted for Word documents using python-docx.

Usage example:
    from services.word_service import create_enhanced_transcript_docx
    
    # Create Word document
    word_buffer = await create_enhanced_transcript_docx(
        title="My Transcript",
        clean_transcript="Clean text without timestamps...",
        full_transcript="[00:00 - 00:05] SPEAKER_1 - Full text with timestamps..."
    )
    
    # Save to file
    with open("transcript.docx", "wb") as f:
        f.write(word_buffer.getvalue())

Features:
    - Title page with Whisper AI logo (without borders) and branding
    - Working table of contents with clickable internal links
    - Clean version section (text only)
    - Full version section (with timestamps and speakers)
    - Professional formatting with consistent styles
    - Automatic fallback if logo file is not found
    - No page breaks (uses spacing instead to avoid display issues)
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from fluentogram import TranslatorRunner


async def create_enhanced_transcript_docx(title: str, clean_transcript: str, full_transcript: str, i18n: TranslatorRunner) -> io.BytesIO:
    """
    Creates an enhanced Word document with title, table of contents, and both clean and full transcript versions.
    Based on the same structure as create_enhanced_transcript_pdf function for consistency.
    
    Args:
        title: Document title
        clean_transcript: Clean version of the transcript (without timestamps/speakers)
        full_transcript: Full version with timestamps and speakers
        i18n: TranslatorRunner instance for translations
        
    Returns:
        io.BytesIO: Buffer with enhanced Word document content
    """
    # Create new document
    doc = Document()
    
    # ===== CONFIGURE STYLES =====
    _configure_document_styles(doc)
    
    # ===== TITLE PAGE =====
    _add_title_page(doc, title, i18n)
    
    # ===== TABLE OF CONTENTS =====
    _add_table_of_contents(doc, i18n)
    
    # Add extra spacing before content sections instead of page break
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(30)
    
    # ===== CLEAN VERSION =====
    _add_clean_version_section(doc, clean_transcript, i18n)
    
    # Add extra spacing before full version instead of page break
    spacing_para2 = doc.add_paragraph()
    spacing_para2.paragraph_format.space_after = Pt(30)
    
    # ===== FULL VERSION =====
    _add_full_version_section(doc, full_transcript, i18n)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


async def create_simple_transcript_docx(transcript: str, title: str, i18n: TranslatorRunner, transcript_type: str = 'clean') -> io.BytesIO:
    """
    Creates a beautifully formatted Word document with a single transcript (keeping all original styling).
    Based on create_enhanced_transcript_docx but without table of contents and dual versions.
    
    Args:
        transcript: The transcript text
        title: Document title
        i18n: TranslatorRunner instance for translations
        transcript_type: Type of transcript ('clean' or 'full')
        
    Returns:
        io.BytesIO: Buffer with enhanced Word document content
    """
    # Create new document
    doc = Document()
    
    # ===== CONFIGURE STYLES =====
    _configure_document_styles(doc)
    
    # ===== TITLE PAGE WITH BRANDING =====
    # Add logo and Whisper AI attribution at the top
    attribution_para = doc.add_paragraph()
    attribution_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    attribution_para.paragraph_format.space_after = Pt(20)

    # Add "Сделано в " in gray
    made_with_text = i18n.transcription_file_made_with_prefix() + ' '
    run1 = attribution_para.add_run(made_with_text)
    run1.font.name = 'Arial'
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(100, 100, 100)

    # Add "Whisper AI" as hyperlink in blue
    whisper_ai_text = i18n.transcription_file_whisper_ai_text()
    run2 = attribution_para.add_run(whisper_ai_text)
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 0, 255)
    _add_hyperlink(attribution_para, run2, 'https://t.me/WhisperSummaryAI_bot', whisper_ai_text)
    
    # Add document title
    title_para = doc.add_paragraph(title)
    title_para.style = 'CustomTitle'
    title_para.paragraph_format.space_after = Pt(15)
    
    # Add version and date in one line
    current_date = datetime.now().strftime("%d.%m.%Y")
    version_title = i18n.google_docs_clean_version_title() if transcript_type == 'clean' else i18n.google_docs_full_version_title()
    version_date_text = f"{version_title}. {i18n.google_docs_creation_date(date=current_date)}"
    
    date_para = doc.add_paragraph(version_date_text)
    date_para.style = doc.styles['Normal']
    date_para.runs[0].font.name = 'Arial'
    date_para.runs[0].font.size = Pt(12)
    date_para.paragraph_format.space_after = Pt(30)
    
    # ===== TRANSCRIPT CONTENT =====
    # Add transcript content with proper styling
    content_para = doc.add_paragraph(transcript)
    content_para.style = 'Content'
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def _configure_document_styles(doc: Document):
    """Configure custom styles for the document."""
    styles = doc.styles
    
    # Create title style
    if 'CustomTitle' not in [style.name for style in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(22)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_style.paragraph_format.space_after = Pt(12)
    
    # Create section header style
    if 'SectionHeader' not in [style.name for style in styles]:
        section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Arial'
        section_style.font.size = Pt(18)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 0, 0)
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        section_style.paragraph_format.space_before = Pt(30)  # Increased from 12
        section_style.paragraph_format.space_after = Pt(15)   # Increased from 6
    
    # Create description style
    if 'Description' not in [style.name for style in styles]:
        desc_style = styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
        desc_style.font.name = 'Arial'
        desc_style.font.size = Pt(11)
        desc_style.font.color.rgb = RGBColor(80, 80, 80)
        desc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        desc_style.paragraph_format.space_after = Pt(8)
    
    # Create content style
    if 'Content' not in [style.name for style in styles]:
        content_style = styles.add_style('Content', WD_STYLE_TYPE.PARAGRAPH)
        content_style.font.name = 'Arial'
        content_style.font.size = Pt(14)
        content_style.font.color.rgb = RGBColor(0, 0, 0)
        content_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        content_style.paragraph_format.line_spacing = 1.2
    
    # Create TOC style
    if 'TOCEntry' not in [style.name for style in styles]:
        toc_style = styles.add_style('TOCEntry', WD_STYLE_TYPE.PARAGRAPH)
        toc_style.font.name = 'Arial'
        toc_style.font.size = Pt(14)
        toc_style.font.color.rgb = RGBColor(0, 0, 255)  # Blue for links
        toc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_style.paragraph_format.space_after = Pt(6)


def _add_title_page(doc: Document, title: str, i18n: TranslatorRunner):
    """Add title page with Whisper AI branding and document info."""
    
    # Add logo and Whisper AI attribution at the top
    attribution_para = doc.add_paragraph()
    attribution_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    attribution_para.paragraph_format.space_after = Pt(20)
    

    # If logo file is not found, add text only
    attribution_para = doc.add_paragraph()
    attribution_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    attribution_para.paragraph_format.space_after = Pt(20)

    # Add "Сделано в " in gray
    made_with_text = i18n.transcription_file_made_with_prefix() + ' '
    run1 = attribution_para.add_run(made_with_text)
    run1.font.name = 'Arial'
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(100, 100, 100)

    # Add "Whisper AI" as hyperlink in blue
    whisper_ai_text = i18n.transcription_file_whisper_ai_text()
    run2 = attribution_para.add_run(whisper_ai_text)
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 0, 255)
    _add_hyperlink(attribution_para, run2, 'https://t.me/WhisperSummaryAI_bot', whisper_ai_text)
    
    # Add document title
    title_para = doc.add_paragraph(title)
    title_para.style = 'CustomTitle'
    title_para.paragraph_format.space_after = Pt(15)
    
    # Add creation date
    current_date = datetime.now().strftime("%d.%m.%Y")
    date_para = doc.add_paragraph(i18n.transcription_file_date_pdf(date=current_date))
    date_para.style = doc.styles['Normal']
    date_para.runs[0].font.name = 'Arial'
    date_para.runs[0].font.size = Pt(12)
    date_para.paragraph_format.space_after = Pt(20)


def _add_table_of_contents(doc: Document, i18n: TranslatorRunner):
    """Add table of contents section."""
    
    # Add TOC header
    toc_header = doc.add_paragraph(i18n.transcription_file_table_of_contents_title())
    toc_header.style = 'SectionHeader'
    toc_header.paragraph_format.space_after = Pt(10)
    
    # Add TOC entries with working hyperlinks
    clean_entry = doc.add_paragraph()
    clean_entry.style = 'TOCEntry'
    _add_internal_hyperlink(clean_entry, i18n.transcription_file_toc_clean_version(), 'clean_version')
    
    full_entry = doc.add_paragraph()
    full_entry.style = 'TOCEntry'
    _add_internal_hyperlink(full_entry, i18n.transcription_file_toc_full_version(), 'full_version')


def _add_clean_version_section(doc: Document, clean_transcript: str, i18n: TranslatorRunner):
    """Add clean version section with description and content."""
    
    # Section header with bookmark
    header = doc.add_paragraph(i18n.transcription_file_clean_version_title())
    header.style = 'SectionHeader'
    _add_bookmark(header, 'clean_version')
    
    # Section description
    clean_description = i18n.transcription_file_clean_version_desc_pdf()
    
    desc_para = doc.add_paragraph(clean_description)
    desc_para.style = 'Description'
    
    # Content
    content_para = doc.add_paragraph(clean_transcript)
    content_para.style = 'Content'


def _add_full_version_section(doc: Document, full_transcript: str, i18n: TranslatorRunner):
    """Add full version section with description and content."""
    
    # Section header with bookmark
    header = doc.add_paragraph(i18n.transcription_file_full_version_title())
    header.style = 'SectionHeader'
    _add_bookmark(header, 'full_version')
    
    # Section description
    full_description = i18n.transcription_file_full_version_desc_pdf()
    
    desc_para = doc.add_paragraph(full_description)
    desc_para.style = 'Description'
    
    # Content
    content_para = doc.add_paragraph(full_transcript)
    content_para.style = 'Content'


def _remove_image_border_completely(picture):
    """Completely remove all borders from an image using a simplified approach."""
    try:
        # Get the inline element
        inline = picture._inline
        
        # Set drawing properties to remove any margins/borders
        drawing_element = inline.get_or_add_cNvPr()
        
        # Get the graphic element and navigate to picture
        graphic = inline.graphic
        graphic_data = graphic.graphicData
        
        # Find namespaces
        a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        pic_ns = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        
        # Find the picture element
        pic_element = graphic_data.find('.//{%s}pic' % pic_ns)
        
        if pic_element is not None:
            # Find all spPr elements and modify them
            for sp_pr in pic_element.findall('.//{%s}spPr' % a_ns):
                # Remove existing line elements that might create borders
                for line in sp_pr.findall('.//{%s}ln' % a_ns):
                    sp_pr.remove(line)
                
                # Add a line element that explicitly has no fill
                line_element = parse_xml(
                    f'<a:ln xmlns:a="{a_ns}" w="0"><a:noFill/></a:ln>'
                )
                sp_pr.insert(0, line_element)
                
    except Exception:
        # Silently continue if border removal fails
        pass


def _add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    # Generate unique ID for bookmark
    import random
    bookmark_id = str(random.randint(1000, 9999))
    
    # Create bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    
    # Insert bookmark start at the beginning of the paragraph
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)


def _add_internal_hyperlink(paragraph, text, bookmark_name):
    """Add an internal hyperlink that links to a bookmark."""
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # Create run with hyperlink formatting
    run = OxmlElement('w:r')
    
    # Add run properties (blue color and underline)
    run_props = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    
    # Font settings
    font = OxmlElement('w:rFonts')
    font.set(qn('w:ascii'), 'Arial')
    font.set(qn('w:hAnsi'), 'Arial')
    
    font_size = OxmlElement('w:sz')
    font_size.set(qn('w:val'), '28')  # 14pt * 2 (Word uses half-points)
    
    run_props.append(font)
    run_props.append(font_size)
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    
    # Add hyperlink to paragraph
    paragraph._element.append(hyperlink)


def _add_hyperlink(paragraph, run, url, text):
    """
    Add a hyperlink to a paragraph run.
    This is a workaround since python-docx doesn't have native hyperlink support.
    """
    # Get the run's rPr (run properties) element
    r_pr = run._element.get_or_add_rPr()
    
    # Create hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create new run for hyperlink
    new_run = OxmlElement('w:r')
    
    # Copy run properties
    r_pr_copy = OxmlElement('w:rPr')
    r_pr_copy.append(OxmlElement('w:color'))
    r_pr_copy.find(qn('w:color')).set(qn('w:val'), '0000FF')  # Blue color
    r_pr_copy.append(OxmlElement('w:u'))
    r_pr_copy.find(qn('w:u')).set(qn('w:val'), 'single')  # Underline
    new_run.append(r_pr_copy)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    
    # Replace the original run with hyperlink
    run._element.getparent().replace(run._element, hyperlink) 